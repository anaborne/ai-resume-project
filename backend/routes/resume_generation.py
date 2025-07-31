from flask import Blueprint, request, jsonify, send_file, current_app
from services.openai_service import generate_optimized_resume
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import subprocess

resume_generation_bp = Blueprint('resume_generation', __name__)

# generates raw optimized resume as JSON
@resume_generation_bp.route('/api/generate-resume', methods=['POST'])
def generate_resume():
    job_description = request.json.get('jobDescription')
    resume_filename = request.json.get('resumeFilename')

    if not resume_filename:
        return jsonify({"error": "Missing resume filename"}), 400
    
    resume_path = os.path.join(current_app.config['UPLOAD_FOLDER'], resume_filename)
    metadata = extract_docx_text(resume_path)
    processed_resume = format_metadata_as_text(metadata)
    optimized_text = generate_optimized_resume(job_description, processed_resume)
    parsed_metadata = parse_llm_formatted_output(optimized_text)
    write_to_docx(parsed_metadata, 'outputs/optimized_resume.docx')

    return jsonify({"optimizedResume": optimized_text}), 200

# convert template docx to text w/ metadata
def extract_docx_text(file_path):
    doc = Document(file_path)
    formatted_output = []

    for i, p in enumerate(doc.paragraphs):
        para_info = {
            "paragraph_number": i + 1,
            "style": p.style.name if p.style else "Unknown",
            "alignment": str(p.alignment),
            "runs": []
        }

        for run in p.runs:
            run_info = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None
            }
            para_info["runs"].append(run_info)
        formatted_output.append(para_info)
    print(formatted_output)
    return formatted_output

# convert template docx parsing into readable text
def format_metadata_as_text(metadata):
    lines = []

    for p in metadata:
        lines.append(f"[PARAGRAPH {p['paragraph_number']}]")
        lines.append(f"STYLE: {p['style']}")
        lines.append(f"ALIGNMENT: {p['alignment']}")
        lines.append("RUNS:")

        for run in p['runs']:
            lines.append(f'  - TEXT: "{run["text"]}"')
            if run.get("bold") is not None:
                lines.append(f"    BOLD: {bool(run['bold'])}")
            if run.get("italic") is not None:
                lines.append(f"    ITALIC: {bool(run['italic'])}")
            if run.get("underline") is not None:
                lines.append(f"    UNDERLINE: {bool(run['underline'])}")
            if run.get("font_size") is not None:
                lines.append(f"    FONT_SIZE: {run['font_size']}")
            if run.get("font_name") is not None:
                lines.append(f"    FONT_NAME: \"{run['font_name']}\"")

        lines.append("")  # spacer line between paragraphs

    return "\n".join(lines)

# parse the LLM output back into metadata format
def parse_llm_formatted_output(formatted_text):
    lines = formatted_text.strip().split('\n')
    parsed_metadata = []
    i = 0

    while i < len(lines):
        if lines[i].startswith('[PARAGRAPH'):
            para_info = {}
            para_info['paragraph_number'] = int(lines[i].split()[1].strip(']'))

            # Next lines are STYLE and ALIGNMENT
            para_info['style'] = lines[i + 1].replace('STYLE: ', '').strip()
            para_info['alignment'] = lines[i + 2].replace('ALIGNMENT: ', '').strip()
            para_info['runs'] = []

            # Dynamically find the RUNS: section
            runs_start = i + 3
            while runs_start < len(lines) and lines[runs_start].strip() != 'RUNS:':
                runs_start += 1

            i = runs_start + 1  # move past 'RUNS:'

            current_run = {}
            while i < len(lines) and not lines[i].startswith('[PARAGRAPH'):
                line = lines[i].strip()

                if line.startswith('- TEXT:'):
                    if current_run:  # save previous run if exists
                        para_info['runs'].append(current_run)
                        current_run = {}
                    current_run['text'] = line.replace('- TEXT:', '').strip().strip('"')
                elif 'BOLD' in line:
                    current_run['bold'] = line.split(':')[1].strip().lower() == 'true'
                elif 'ITALIC' in line:
                    current_run['italic'] = line.split(':')[1].strip().lower() == 'true'
                elif 'UNDERLINE' in line:
                    current_run['underline'] = line.split(':')[1].strip().lower() == 'true'
                elif 'FONT_SIZE' in line:
                    try:
                        current_run['font_size'] = float(line.split(':')[1].strip())
                    except:
                        current_run['font_size'] = None
                elif 'FONT_NAME' in line:
                    current_run['font_name'] = line.split(':')[1].strip().strip('"')

                i += 1

            if current_run:
                para_info['runs'].append(current_run)

            parsed_metadata.append(para_info)
        else:
            i += 1

    return parsed_metadata

# write the parsed metadata back to a new docx file
def write_to_docx(parsed_metadata, output_path):
    doc = Document()

    for para in parsed_metadata:
        paragraph = doc.add_paragraph()
        
        alignment = para["alignment"].upper()
        if "CENTER" in alignment:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "RIGHT" in alignment:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for run_meta in para["runs"]:
            run = paragraph.add_run(run_meta["text"])
            
            run.bold = run_meta.get("bold", False)
            run.italic = run_meta.get("italic", False)
            run.underline = run_meta.get("underline", False)
            
            font = run.font
            if run_meta.get("font_size"):
                font.size = Pt(run_meta["font_size"])
            if run_meta.get("font_name"):
                font.name = run_meta["font_name"]
    
    doc.save(output_path)

# Download route for the optimized resume
@resume_generation_bp.route('/api/download-resume', methods=['GET'])
def download_resume():
    output_path = 'outputs/optimized_resume.docx'
    return send_file(output_path, as_attachment=True)

@resume_generation_bp.route('/api/download-resume-pdf', methods=['GET'])
def download_resume_pdf():
    input_path = os.path.join(current_app.config['OUTPUT_FOLDER'], 'optimized_resume.docx')
    output_path = os.path.join(current_app.config['OUTPUT_FOLDER'], 'optimized_resume.pdf')

    subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
        current_app.config['OUTPUT_FOLDER'], input_path
    ])

    return send_file(output_path, mimetype='application/pdf')
